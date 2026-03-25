import logging
import os
from docx import Document


class Exporter:
    """
    Exports Markdown content into Word Document templates via python-docx & Jinja2.
    """

    def __init__(self, template_dir="templates"):
        self.template_dir = template_dir
        if not os.path.exists(self.template_dir):
            os.makedirs(self.template_dir)

        self.env = None

    def export_word(self, exam_data: dict, template_name: str, output_path: str):
        """
        Takes purely markdown/html content and injects into Word template placeholders.
        """
        try:
            template_path = os.path.normpath(
                os.path.join(self.template_dir, os.path.basename(template_name))
            )
            if not os.path.exists(template_path):
                # Create dummy template if missing
                doc = Document()
                doc.add_heading("{{ school_name }} Exam", 0)
                doc.add_paragraph("{{ content }}")
                doc.save(template_path)

            # Using python-docx for native parsing instead of simple Jinja due to OOXML format
            # This is a simplified placeholder for the actual complex HTML2Word conversion.
            doc = Document(template_path)

            for p in doc.paragraphs:
                if "{{ school_name }}" in p.text:
                    p.text = p.text.replace(
                        "{{ school_name }}", exam_data.get("school", "SmartQB Academy")
                    )
                if "{{ content }}" in p.text:
                    # In real app, replace with converted Markdown HTML blocks
                    p.text = p.text.replace(
                        "{{ content }}",
                        exam_data.get("markdown", "Questions go here..."),
                    )

            doc.save(output_path)
            return True
        except Exception:
            logging.exception("Export failed in export function")
            return False
